from mcp.server.fastmcp import FastMCP

mcp = FastMCP(
    'Custom Tools'
)

import json
import os
from docx import Document
from docx.shared import Pt

@mcp.tool(
    name="create_docx",
    description="Create a Word document (.docx) with specified text, headings, and formatting"
)
def create_docx_tool(input_data, file_path=None):
    """
    Create a .docx file with structured content.

    Args:
        input_data: JSON string, dictionary, or plain text string containing document content.
        file_path: Optional file path for plain text input (required if input_data is a plain text string).

    Expected input format (JSON string or dictionary):
    {
        "file_path": "path/to/output.docx",           # Required; path to save the document
        "title": "Document Title",                   # Optional; main document title
        "text": [                                    # Optional; list of content items
            {"type": "heading", "text": "Heading Text", "level": 1},  # Heading with optional level (1-9)
            {"type": "paragraph", "text": "Sample Paragraph"},        # Paragraph text
            {"type": "bullet_list", "items": ["Item 1", "Item 2"]},   # Bullet list
            {"type": "numbered_list", "items": ["Item 1", "Item 2"]}  # Numbered list
        ]
    }
    OR plain text string (treated as a single paragraph, requires file_path parameter)

    Returns:
        dict: {
            "status": "success" or "error",
            "message": Descriptive message,
            "file_path": Path to the created file (on success)
        }

    Raises:
        ValueError: If input is invalid (e.g., missing file_path, invalid JSON, unsupported type).
        IOError: If the file cannot be saved (e.g., invalid path or permissions issue).

    Examples:
        JSON input:
            create_docx_tool('{
                "file_path": "output.docx",
                "title": "My Document",
                "text": [
                    {"type": "heading", "text": "Section 1", "level": 1},
                    {"type": "paragraph", "text": "This is a paragraph."},
                    {"type": "bullet_list", "items": ["Item 1", "Item 2"]}
                ]
            }')

        Plain text input:
            create_docx_tool("This is a simple document.", file_path="output.docx")
    """
    try:
        # Parse input if it's a JSON string
        if isinstance(input_data, str):
            try:
                data = json.loads(input_data)
            except json.JSONDecodeError:
                # Handle plain text input
                if not file_path:
                    raise ValueError("file_path is required for plain text input")
                data = {"file_path": file_path, "text": [{"type": "paragraph", "text": input_data}]}
        else:
            data = input_data

        # Validate input type
        if not isinstance(data, dict):
            raise ValueError("Input must be a JSON string or dictionary")

        # Validate file_path
        if "file_path" not in data:
            raise ValueError("file_path is required in input data")
        file_path = data["file_path"]
        if not isinstance(file_path, str) or not file_path.strip():
            raise ValueError("file_path must be a non-empty string")
        if not file_path.lower().endswith(".docx"):
            file_path += ".docx"

        # Create a new Document
        doc = Document()
        content_added = False

        # Add title if provided
        if "title" in data and data["title"]:
            if not isinstance(data["title"], str):
                raise ValueError("title must be a string")
            title = doc.add_heading(data["title"], 0)
            title.style.font.size = Pt(16)  # Ensure consistent title formatting
            content_added = True

        # Process text content
        if "text" in data:
            if not isinstance(data["text"], list):
                raise ValueError("text must be a list")
            for item in data["text"]:
                if not isinstance(item, dict) or "type" not in item:
                    raise ValueError("Each text item must be a dictionary with a 'type' field")
                
                item_type = item["type"]
                if item_type == "heading":
                    if "text" not in item:
                        raise ValueError("Heading must have a 'text' field")
                    level = item.get("level", 1)
                    if not isinstance(level, int) or level < 1 or level > 9:
                        level = 1  # Default to level 1 if invalid
                    doc.add_heading(item["text"], level)
                    content_added = True
                
                elif item_type == "paragraph":
                    if "text" not in item:
                        raise ValueError("Paragraph must have a 'text' field")
                    doc.add_paragraph(item["text"])
                    content_added = True
                
                elif item_type == "bullet_list":
                    if "items" not in item or not isinstance(item["items"], list):
                        raise ValueError("Bullet list must have an 'items' list")
                    for bullet_item in item["items"]:
                        doc.add_paragraph(str(bullet_item), style='List Bullet')
                    content_added = True
                
                elif item_type == "numbered_list":
                    if "items" not in item or not isinstance(item["items"], list):
                        raise ValueError("Numbered list must have an 'items' list")
                    for num_item in item["items"]:
                        doc.add_paragraph(str(num_item), style='List Number')
                    content_added = True
                
                else:
                    raise ValueError(f"Unsupported item type: {item_type}")

        # Ensure the document is not empty
        if not content_added:
            raise ValueError("No content provided; document would be empty")

        # Create directory if it doesn't exist
        os.makedirs(os.path.dirname(file_path) or ".", exist_ok=True)

        # Save the document
        doc.save(file_path)

        return {
            "status": "success",
            "message": f"Document created successfully at {file_path}",
            "file_path": file_path
        }

    except json.JSONDecodeError as e:
        return {
            "status": "error",
            "message": f"Invalid JSON input: {str(e)}"
        }
    except ValueError as e:
        return {
            "status": "error",
            "message": f"Input validation error: {str(e)}"
        }
    except IOError as e:
        return {
            "status": "error",
            "message": f"Failed to save document: {str(e)}"
        }
    except Exception as e:
        return {
            "status": "error",
            "message": f"Unexpected error: {str(e)}"
        }

import json
import os
from docx import Document

@mcp.tool(
    name="read_docx",
    description="Read content from an existing Word document (.docx)"
)
def read_docx_tool(input_data):
    """
    Read content from an existing .docx file and return it in a structured format.

    Args:
        input_data: JSON string, dictionary, or plain string specifying the .docx file path.

    Expected input format (JSON string or dictionary):
    {
        "file_path": "path/to/existing.docx"  # Required; the .docx file to read
    }
    OR plain string (treated as file_path)

    Returns:
        dict: {
            "status": "success" or "error",
            "message": Descriptive message,
            "content": [  # On success; list of content items
                {"type": "heading", "text": "Heading Text", "level": int},
                {"type": "paragraph", "text": "Sample Paragraph"},
                {"type": "bullet_list", "items": ["Item 1", "Item 2"]},
                {"type": "numbered_list", "items": ["Step 1", "Step 2"]}
            ]
        }

    Raises:
        ValueError: If input is invalid (e.g., missing or invalid file_path).
        FileNotFoundError: If the .docx file does not exist.
        IOError: If the file cannot be read.

    Examples:
        JSON input:
            read_docx_tool('{"file_path": "existing.docx"}')

        Plain string input:
            read_docx_tool("existing.docx")
    """
    try:
        # Parse input
        if isinstance(input_data, str):
            try:
                data = json.loads(input_data)
            except json.JSONDecodeError:
                data = {"file_path": input_data}
        else:
            data = input_data

        # Validate input type
        if not isinstance(data, dict):
            raise ValueError("Input must be a JSON string, dictionary, or plain string")

        # Validate file_path
        if "file_path" not in data:
            raise ValueError("file_path is required in input data")
        file_path = data["file_path"]
        if not isinstance(file_path, str) or not file_path.strip():
            raise ValueError("file_path must be a non-empty string")
        if not file_path.lower().endswith(".docx"):
            file_path += ".docx"

        # Check if file exists
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Input file {file_path} does not exist")

        # Load the document
        doc = Document(file_path)
        content = []

        # Track list items for bullet/numbered lists
        current_list = None
        current_list_type = None

        # Read paragraphs
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                continue  # Skip empty paragraphs

            style = paragraph.style.name if paragraph.style else "Normal"

            # Detect headings (e.g., Heading 1, Heading 2)
            if style.startswith("Heading "):
                try:
                    level = int(style.split(" ")[1])
                    content.append({"type": "heading", "text": paragraph.text, "level": level})
                    current_list = None  # Reset list tracking
                except ValueError:
                    content.append({"type": "paragraph", "text": paragraph.text})
            
            # Detect bullet lists
            elif style == "List Bullet":
                if current_list_type != "bullet_list":
                    current_list = {"type": "bullet_list", "items": []}
                    content.append(current_list)
                    current_list_type = "bullet_list"
                current_list["items"].append(paragraph.text)
            
            # Detect numbered lists
            elif style == "List Number":
                if current_list_type != "numbered_list":
                    current_list = {"type": "numbered_list", "items": []}
                    content.append(current_list)
                    current_list_type = "numbered_list"
                current_list["items"].append(paragraph.text)
            
            # Treat other styles as paragraphs
            else:
                content.append({"type": "paragraph", "text": paragraph.text})
                current_list = None  # Reset list tracking

        return {
            "status": "success",
            "message": f"Document read successfully from {file_path}",
            "content": content
        }

    except json.JSONDecodeError as e:
        return {
            "status": "error",
            "message": f"Invalid JSON input: {str(e)}"
        }
    except ValueError as e:
        return {
            "status": "error",
            "message": f"Input validation error: {str(e)}"
        }
    except FileNotFoundError as e:
        return {
            "status": "error",
            "message": f"File error: {str(e)}"
        }
    except IOError as e:
        return {
            "status": "error",
            "message": f"Failed to read document: {str(e)}"
        }
    except Exception as e:
        return {
            "status": "error",
            "message": f"Unexpected error: {str(e)}"
        }


import json
import os
from docx import Document
from docx.shared import Pt

@mcp.tool(
    name="edit_docx",
    description="Edit an existing Word document (.docx) by appending or Replacing content"
)
def edit_docx_tool(input_data, file_path=None, output_path=None):
    """
    Edit an existing .docx file by appending content or replacing text.

    Args:
        input_data: JSON string, dictionary, or plain text string containing edit instructions.
        file_path: Optional; path to the existing .docx file (required for plain text input).
        output_path: Optional; path to save the modified file (defaults to file_path).

    Expected input format (JSON string or dictionary):
    {
        "file_path": "path/to/existing.docx",         # Required; existing .docx file
        "output_path": "path/to/output.docx",         # Optional; defaults to file_path
        "operation": "append" or "replace",           # Required; type of edit
        "text": [                                    # Optional; for append operation
            {"type": "heading", "text": "Heading Text", "level": 1},
            {"type": "paragraph", "text": "Sample Paragraph"},
            {"type": "bullet_list", "items": ["Item 1", "Item 2"]},
            {"type": "numbered_list", "items": ["Step 1", "Step 2"]}
        ],
        "replace": {                                 # Optional; for replace operation
            "find": "text to find",
            "replace_with": "new text"
        }
    }
    OR plain text string (treated as a paragraph to append, requires file_path)

    Returns:
        dict: {
            "status": "success" or "error",
            "message": Descriptive message,
            "file_path": Path to the modified file (on success)
        }

    Raises:
        ValueError: If input is invalid (e.g., missing file_path, invalid operation).
        IOError: If the file cannot be loaded or saved.
        FileNotFoundError: If the input .docx file does not exist.

    Examples:
        Append JSON input:
            edit_docx_tool('{
                "file_path": "existing.docx",
                "output_path": "modified.docx",
                "operation": "append",
                "text": [
                    {"type": "heading", "text": "New Section", "level": 1},
                    {"type": "paragraph", "text": "Added text."}
                ]
            }')

        Replace JSON input:
            edit_docx_tool('{
                "file_path": "existing.docx",
                "operation": "replace",
                "replace": {"find": "old", "replace_with": "new"}
            }')

        Plain text input (append):
            edit_docx_tool("New note!", file_path="existing.docx", output_path="modified.docx")
    """
    try:
        # Parse input if it's a JSON string
        if isinstance(input_data, str):
            try:
                data = json.loads(input_data)
            except json.JSONDecodeError:
                if not file_path:
                    raise ValueError("file_path is required for plain text input")
                data = {
                    "file_path": file_path,
                    "output_path": output_path or file_path,
                    "operation": "append",
                    "text": [{"type": "paragraph", "text": input_data}]
                }
        else:
            data = input_data

        # Validate input type
        if not isinstance(data, dict):
            raise ValueError("Input must be a JSON string or dictionary")

        # Validate file_path
        if "file_path" not in data:
            raise ValueError("file_path is required in input data")
        file_path = data["file_path"]
        if not isinstance(file_path, str) or not file_path.strip():
            raise ValueError("file_path must be a non-empty string")
        if not file_path.lower().endswith(".docx"):
            file_path += ".docx"

        # Validate output_path
        output_path = data.get("output_path", file_path)
        if not isinstance(output_path, str) or not output_path.strip():
            raise ValueError("output_path must be a non-empty string")
        if not output_path.lower().endswith(".docx"):
            output_path += ".docx"

        # Validate operation
        if "operation" not in data:
            raise ValueError("operation is required in input data")
        operation = data["operation"]
        if operation not in ["append", "replace"]:
            raise ValueError("operation must be 'append' or 'replace'")

        # Check if file exists
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Input file {file_path} does not exist")

        # Load the existing document
        doc = Document(file_path)
        content_modified = False

        # Perform the edit based on operation
        if operation == "append":
            if "text" not in data:
                raise ValueError("text is required for append operation")
            if not isinstance(data["text"], list):
                raise ValueError("text must be a list")
            for item in data["text"]:
                if not isinstance(item, dict) or "type" not in item:
                    raise ValueError("Each text item must be a dictionary with a 'type' field")
                
                item_type = item["type"]
                if item_type == "heading":
                    if "text" not in item:
                        raise ValueError("Heading must have a 'text' field")
                    level = item.get("level", 1)
                    if not isinstance(level, int) or level < 1 or level > 9:
                        level = 1
                    doc.add_heading(item["text"], level)
                    content_modified = True
                
                elif item_type == "paragraph":
                    if "text" not in item:
                        raise ValueError("Paragraph must have a 'text' field")
                    doc.add_paragraph(item["text"])
                    content_modified = True
                
                elif item_type == "bullet_list":
                    if "items" not in item or not isinstance(item["items"], list):
                        raise ValueError("Bullet list must have an 'items' list")
                    for bullet_item in item["items"]:
                        doc.add_paragraph(str(bullet_item), style='List Bullet')
                    content_modified = True
                
                elif item_type == "numbered_list":
                    if "items" not in item or not isinstance(item["items"], list):
                        raise ValueError("Numbered list must have an 'items' list")
                    for num_item in item["items"]:
                        doc.add_paragraph(str(num_item), style='List Number')
                    content_modified = True
                
                else:
                    raise ValueError(f"Unsupported item type: {item_type}")

        elif operation == "replace":
            if "replace" not in data or not isinstance(data["replace"], dict):
                raise ValueError("replace must be a dictionary for replace operation")
            if "find" not in data["replace"] or "replace_with" not in data["replace"]:
                raise ValueError("replace must include 'find' and 'replace_with' fields")
            
            find_text = str(data["replace"]["find"])
            replace_text = str(data["replace"]["replace_with"])
            
            for paragraph in doc.paragraphs:
                if find_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(find_text, replace_text)
                    content_modified = True

        # Ensure some modification was made
        if not content_modified:
            raise ValueError("No changes were made to the document")

        # Create directory for output if it doesn't exist
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

        # Save the document
        doc.save(output_path)

        return {
            "status": "success",
            "message": f"Document modified successfully at {output_path}",
            "file_path": output_path
        }

    except json.JSONDecodeError as e:
        return {
            "status": "error",
            "message": f"Invalid JSON input: {str(e)}"
        }
    except ValueError as e:
        return {
            "status": "error",
            "message": f"Input validation error: {str(e)}"
        }
    except FileNotFoundError as e:
        return {
            "status": "error",
            "message": f"File error: {str(e)}"
        }
    except IOError as e:
        return {
            "status": "error",
            "message": f"Failed to load or save document: {str(e)}"
        }
    except Exception as e:
        return {
            "status": "error",
            "message": f"Unexpected error: {str(e)}"
        }


import os

@mcp.tool(
        name="delete_file",
        description="Delete a file at a given path"
)
def delete_file_tool(input):
    file_path = input
    input = "C:\\testfolder"
    if os.path.exists(file_path):
        os.remove(file_path)
        return {"status": "deleted"}
    else:
        return {"status": "file not found"}
    

#Death Tool Trial 1
import fitz  # PyMuPDF

@mcp.tool(
    name="Extract_Document_Metadata",
    description="Extracts raw text from files and asks Claude to identify key metadata fields and their values for structured document understanding."
)
def extract_death_info_tool(file_path: str) -> dict:
    try:
        # Step 1: Read text from the PDF
        doc = fitz.open(file_path)
        full_text = ""
        for page in doc:
            full_text += page.get_text()
        text = full_text.strip()

        DEATH_QUESTIONS = {
    "name_of_deceased": "What is the name of the deceased person in the file provided?",
    "father_name": "What is the Father's name or Mother's name of the deceased person?",
    "date_of_accident": "What is the date particulars of the accident given in the text?",
    "place_of_accident": "What is the place of accident given?",
    "number_of_dependents": "What is the total number of dependents/dependants of the deceased person, including their relationship with the deceased? Please provide the count as a numerical value, followed by the word 'dependents' or 'dependants' and the relationship of the dependents or dependants with the deceased (e.g., '3 dependents: wife, 2 children').",
    "age_of_deceased": "Age of deceased at the time of accident?",
    "annual_income": "What is the annual income of the deceased person? "
                    "If the text contains the word 'monthly' or 'per month' or 'PM' calculate the annual income "
                    "by multiplying the monthly income by 12 and provide the final integer value "
                    "in the format 'X,XXX' (without quotes). "
                    "If the text contains the phrase 'annual income' or 'per year' or 'per annum', "
                    "provide the annual income as the final integer value in the format 'X,XXX' (without quotes). "
                    "If neither the annual nor the monthly income is available, consider the value as 0.",
    "occupation": "occupation or designation of the deceased person?",
    "medical_expenses": "What is the medical expense/medical bills/amount of expense on treatment involved with deceased person?",
    "tribunal_awarded":"what is the compensation amount awarded by the tribunal in the case file?"
    }


        # Step 2: Prepare prompt for Claude
        questions_text = "\n".join([f"{key}: {q}" for key, q in DEATH_QUESTIONS.items()])
        prompt = f"""{text}

Answer the following questions based on the above data and return the response in JSON format. Keep the exact questions as the keys:
{questions_text}

Your response should be a valid JSON object with no explanations before or after.
"""

        return {
            "status": "success",
            "prompt": prompt
        }

    except Exception as e:
        return {
            "status": "error",
            "message": str(e)
        }

#General Metadata Extraction Tool Trial 1
import fitz  # PyMuPDF

class ToolException(Exception):
    """Custom exception for tool-specific errors."""
    pass

from typing import Dict, Generator 

@mcp.tool(
    name="Extract_Document_Metadata",
    description="Extracts text from any document and uses Claude to identify key metadata fields and their corresponding values."
)
def extract_document_metadata(file_path: str) -> Generator[Dict, Dict, Dict]:
    """
    Tool that processes a document file to extract metadata.
    Steps:
    1. Extract raw text from file
    2. Ask Claude to identify important metadata keys
    3. Ask Claude to extract corresponding values for those keys
    """

    try:
        # Step 1: Read text from the PDF
        doc = fitz.open(file_path)
        full_text = ""
        for page in doc:
            full_text += page.get_text()
        text = full_text.strip()

        if not full_text or len(full_text.strip()) < 20:
            raise ToolException("The uploaded file appears empty or unreadable.")

        # Optional: Truncate if too large (Claude context safety)
        document_excerpt = full_text[:3000]

        # Step 2: Ask Claude to identify metadata keys
        metadata_keys_prompt = (
            f"You are given the following document:\n\n"
            f"{document_excerpt}\n\n"
            "Your task is to identify all the important metadata keys from this document. "
            "Metadata keys represent essential information fields such as: Title, Case Number, Date, Name, Summary, Organization, etc. "
            "Return only a clean bullet list of metadata keys without values."
        )

        yield {"role": "user", "content": metadata_keys_prompt}

        # Step 3: Once keys are returned, ask Claude to extract corresponding values
        def on_keys_response(response):
            keys_raw = response.get("content", "").strip()
            if not keys_raw:
                raise ToolException("Claude did not return any metadata keys.")

            # Format the prompt to extract values for each key
            metadata_values_prompt = (
                f"The document is:\n\n{document_excerpt}\n\n"
                f"The metadata keys identified are:\n{keys_raw}\n\n"
                "Now extract and return the corresponding value for each key. "
                "Respond in the format:\n- Key: Value\n- Key: Value\n..."
            )

            return {"role": "user", "content": metadata_values_prompt}

        yield on_keys_response

    except Exception as e:
        raise ToolException(f"An error occurred: {str(e)}")


if __name__ == "__main__":  
    mcp.run(transport='stdio')    